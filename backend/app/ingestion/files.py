from __future__ import annotations

from pathlib import Path
from uuid import uuid4

import pandas as pd

from app.config import SQLITE_DIR
from app.analysis.segmenter import segment_text
from app.ingestion.text import chunks_from_text, read_text_file
from app.models.schemas import IngestedChunk, Segment
from app.rag.sql import load_dataframe
from app.store import TableInfo


TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
TABLE_EXTENSIONS = {".csv", ".xlsx", ".xls"}


def detect_file_kind(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TABLE_EXTENSIONS:
        return "table"
    if suffix in PDF_EXTENSIONS:
        return "pdf"
    if suffix in DOCX_EXTENSIONS:
        return "docx"
    if suffix in TEXT_EXTENSIONS:
        return "text"
    return "unknown"


def extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append(f"Page {index}\n{page.extract_text() or ''}")
    return "\n\n".join(pages)


def extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            table_lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(paragraphs + table_lines)


def extract_document_text(path: Path) -> tuple[str, str]:
    kind = detect_file_kind(path)
    if kind == "pdf":
        text = extract_pdf_text(path)
    elif kind == "docx":
        text = extract_docx_text(path)
    else:
        text = read_text_file(path)
    return text, kind


def ingest_document_file(dataset_id: str, input_id: str, path: Path) -> tuple[list[IngestedChunk], list[Segment]]:
    text, kind = extract_document_text(path)
    segments = segment_text(text, path.name, input_id)
    indexable_segments = [segment for segment in segments if segment.rag_module != "sql"]
    chunks: list[IngestedChunk] = []
    for segment in indexable_segments:
        chunks.extend(
            chunks_from_text(
                dataset_id=dataset_id,
                input_id=input_id,
                text=segment.text,
                title=segment.title,
                file_name=path.name,
                metadata={
                    "kind": kind,
                    "segment_id": segment.id,
                    "segment_type": segment.segment_type,
                    "rag_module": segment.rag_module,
                    "section_title": segment.title,
                },
            )
        )
    return chunks, segments


def chunks_for_full_document(dataset_id: str, input_id: str, path: Path) -> list[IngestedChunk]:
    text, kind = extract_document_text(path)
    return chunks_from_text(
        dataset_id=dataset_id,
        input_id=input_id,
        text=text,
        title=path.name,
        file_name=path.name,
        metadata={"kind": kind},
    )


def ingest_table_file(dataset_id: str, input_id: str, path: Path) -> TableInfo:
    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
    else:
        df = pd.read_excel(path)
    table_name = f"table_{input_id.replace('-', '_')[:12]}"
    db_path = SQLITE_DIR / f"{dataset_id}.db"
    load_dataframe(df, db_path, table_name)
    return TableInfo(
        input_id=input_id,
        table_name=table_name,
        db_path=db_path,
        columns=[str(c) for c in df.columns],
        row_count=len(df),
        source_name=path.name,
    )


def make_input_record(path: Path, kind: str) -> dict:
    return {
        "id": str(uuid4()),
        "name": path.name,
        "kind": kind,
        "extension": path.suffix.lower(),
    }
